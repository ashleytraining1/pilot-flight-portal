import streamlit as st
import pandas as pd
from pyairtable import Api
from datetime import datetime, date, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import smtplib
from email.message import EmailMessage

def send_currency_alert(pilot_email, last_date, days_left):
    try:
        msg = EmailMessage()
        msg.set_content(f"Captain, your last flight was on {last_date}. You have {days_left} days remaining to maintain your 21-day currency.")
        msg['Subject'] = "✈️ Currency Reminder: 2 Days Remaining"
        msg['From'] = st.secrets["emails"]["smtp_user"]
        msg['To'] = pilot_email

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(st.secrets["emails"]["smtp_user"], st.secrets["emails"]["smtp_pass"])
            server.send_message(msg)
    except Exception as e:
        print(f"Email error: {e}")

# Initialize df_raw as an empty dataframe so the 'if' check doesn't crash
df_raw = pd.DataFrame()
# --- 1. SETUP & THEME ---
st.set_page_config(page_title="Pilot Flight Portal", layout="wide")
st.markdown("<style>.stDataFrame {font-size: 1.05rem;}</style>", unsafe_allow_html=True)

api_key = st.secrets["connections"]["airtable"]["api_key"]
base_id = st.secrets["connections"]["airtable"]["base_id"]
api = Api(api_key)

table_monthly = api.table(base_id, "Monthly Stats") 
df_stats = pd.DataFrame(table_monthly.all())

# --- 2. THE PRECISION FORMATTER (HH:MM) ---
def universal_formatter(val, col_name):
    if val is None or val == "" or pd.isna(val):
        return "0:00" if any(x in col_name.lower() for x in ["time", "flying", "dual", "tp", "instr"]) else "0"
    try:
        if isinstance(val, str): return val
        num = float(val)
        col_lower = col_name.lower()
        if any(x in col_lower for x in ["ldg", "landings", "approach", "hold", "count", "num"]):
            return str(int(round(num)))
        if num >= 10: total_mins = int(round(num / 60))
        else: total_mins = int(round(num * 60))
        return f"{total_mins // 60}:{total_mins % 60:02d}"
    except: return str(val)

# --- DATA HELPERS ---
def get_clean_val(val):
    """Unpacks Airtable lists and converts to number."""
    if isinstance(val, list):
        val = val[0] if len(val) > 0 else 0
    try:
        return float(val) if val is not None else 0.0
    except:
        return 0.0

def safe_sum(dataframe, target_col):
    """The core sum function that matches column names exactly."""
    if dataframe.empty or target_col not in dataframe.columns: 
        return 0
    return dataframe[target_col].apply(get_clean_val).sum()

def get_twin_logic(dataframe, time_col_name):
    """Filters aircraft for C145A/Y12II and sums specific time column."""
    if dataframe.empty or time_col_name not in dataframe.columns: 
        return 0
    type_col = "AIRCRAFT" 
    if type_col not in dataframe.columns: return 0
    mask = dataframe[type_col].astype(str).str.contains('C145A|Y12II', case=False, na=False)
    return dataframe[mask][time_col_name].apply(get_clean_val).sum()

# --- 3. SIDEBAR LOGIN & NAV ---
st.sidebar.title("👨‍✈️ Pilot Portal")
user_email = st.sidebar.text_input("Login with Email").strip().lower()

# --- LEGAL GATEKEEPER LOGIC 
is_duly_signed = False 
user_record = None
u_fields = {}

if user_email:
    # 1. Fetch the user's record from your Summary Table
    summary_table = api.table(base_id, "tbl0apM9eg1tOEYss")
    user_record = next((r for r in summary_table.all() if user_email in str(r['fields'].values()).lower()), None)

    if user_record:
        u_id = user_record['id']
        u_fields = user_record['fields']
        
        # 2. Check the permanent 'Legal_Accepted' field 
        is_duly_signed = u_fields.get("Legal_Accepted", False)

       
        if not is_duly_signed:
            st.sidebar.warning("⚖️ **Action Required**")
            st.sidebar.write("Please read and accept the Terms & Conditions in the main window to proceed.")
            
            st.header("⚖️ Flight Portal Service Agreement & Privacy Policy")
            st.markdown("""
            ### 1. Purpose of Service
            The Pilot Flight Portal is a digital analytical tool designed to provide automated summaries of flight records. This system is a **supplementary convenience** and does not replace the Master Flight Log or official Kenya Air Force (KAF) record-keeping systems.

            ### 2. Data Sovereignty & Integrity
            * **Pilot Ownership:** All flight data remains the intellectual property of the individual pilot. The portal acts only as a processor of this information.
            * **Data Accuracy:** While the system utilizes precision formatting (HH:MM) to eliminate manual calculation errors, the **Pilot in Command (PIC)** holds the ultimate legal responsibility for the accuracy of totals submitted on official documents (CAT 55 / IR MASILA).
            * **System Discrepancy:** In the event of a conflict between portal totals and the Master Log, the Master Log shall remain the primary authority.

            ### 3. Privacy & Security Commitment
            * **Restricted Access:** Your records are strictly partitioned. No other pilot, officer, or external entity can view your flight metrics.
            * **No Third-Party Sharing:** We do not sell, share, or transmit your service number, flight routes, or performance data to any third party or military administrative body.
            * **Encryption:** Data is fetched via secure API protocols and is handled according to standard data protection practices.

            ### 4. Limitation of Liability
            The Developer shall not be held liable for administrative delays, record discrepancies, or any disciplinary actions arising from the use of generated summaries. By proceeding, you acknowledge that you have cross-verified the system's output against your primary records.

            ### 5. Subscription & Access
            Access is granted on a per-user basis. Unauthorized sharing of login credentials or attempts to reverse-engineer the portal logic will result in immediate termination of access without refund.
            """)
            
            # --- THE CHECKBOX & BUTTON FIX ---
            agree_check = st.checkbox("I have read and accepted the Terms and Conditions")
            
            if agree_check:
                if st.button("I Accept the Terms and Conditions"):
                    summary_table.update(u_id, {"Legal_Accepted": True})
                    st.success("Agreement Signed! Loading Portal...")
                    st.rerun()
            
            st.stop() # This keeps the dashboard hidden until they click the button

# --- 4. DASHBOARD ACCESS ---
if user_email and is_duly_signed:
    # We build the 'Smart Link' by adding the prefill and hide parameters to your URL
    smart_url = (
        f"https://airtable.com/appQT0NaMId6xWAya/shrfdHvw2KuSeobV1"
        f"?prefill_Summary={user_email}"
        f"&hide_Summary=true"
    )
    
    st.sidebar.link_button("➕ Log New Flight", smart_url)
    
    if st.sidebar.button("⚖️ Review Terms & Policies"):
        # This allows them to see the terms again if they choose
        summary_table.update(u_id, {"Legal_Accepted": False})
        st.rerun()
    
    st.sidebar.divider()
    st.sidebar.info("**Support:** For system issues or activation, contact admin via WhatsApp.")
    
    # --- YOUR DASHBOARD CODE (Tabs, Metrics, etc.) STARTS HERE ---

# --- 4. MAIN LOGIC ---
if user_email:
    summary_table = api.table(base_id, "tbl0apM9eg1tOEYss")
    logbook_table = api.table(base_id, "tblZEyrmxB2AV68jS")
    stats_table = api.table(base_id, "tblNzncflTuufPioz")

    # We use .all() to get all rows, then extract the 'fields' dictionary from each row
    raw_records = logbook_table.all()
    df_raw = pd.DataFrame([r['fields'] for r in raw_records]) if raw_records else pd.DataFrame()

    stats_records = stats_table.all()
    # This creates the columns like 'Month' and 'Monthly Narrative' that were missing
    df_stats = pd.DataFrame([r['fields'] for r in stats_records]) if stats_records else pd.DataFrame()
    
    with st.spinner('Authenticating...'):
        all_summaries = summary_table.all()
        user_summary_list = [r['fields'] for r in all_summaries if user_email in str(r['fields'].values()).lower()]

    # --- 5. DATA PREP ---
    raw_logs = logbook_table.all()  # Now it can see the table!
    user_logs = [r['fields'] for r in raw_logs if user_email.lower() in str(r['fields'].values()).lower()]
    df_raw = pd.DataFrame(user_logs)

    if not df_raw.empty:
        df_raw.columns = df_raw.columns.str.strip()
        date_col = "LOGBOOK DATE"
        
        df_raw[date_col] = pd.to_datetime(df_raw[date_col], errors='coerce')
        df_raw = df_raw.sort_values(by=date_col, ascending=False)
        
        # Currency Logic (Tested for your 4-day flight)
        last_flight_date = df_raw[date_col].max().date()
        days_since_flight = (date.today() - last_flight_date).days

        if 19 <= days_since_flight < 21:
            days_left = 21 - days_since_flight
            st.warning(f"⚠️ **Currency Warning:** Last flight was {last_flight_date}. You have {days_left} days left.")
            
            if "email_sent" not in st.session_state:
                send_currency_alert(user_email, last_flight_date, days_left)
                st.session_state.email_sent = True

        # 3. IRT / Recent History Logic
        # Make sure this IRT logic is INSIDE the "if not df_raw.empty" block
        six_months_ago = pd.Timestamp(date.today()) - pd.Timedelta(days=180)
        recent_df = df_raw[df_raw[date_col] >= six_months_ago]
    else:
        recent_df = pd.DataFrame()

    if user_summary_list:
        u = user_summary_list[0]
        
        # --- 5. SUBSCRIPTION LOGIC ---
        is_active = u.get("Subscription_Active", False)
        expiry_val = u.get("Subscription_Expiry")
        expiry_date = None
        days_left = 999 
        
        if expiry_val:
            try:
                if isinstance(expiry_val, str):
                    for fmt in ('%Y-%m-%d', '%d/%m/%Y', '%m/%d/%Y', '%d-%m-%Y'):
                        try:
                            expiry_date = datetime.strptime(expiry_val, fmt).date()
                            break
                        except ValueError: continue
                else: expiry_date = expiry_val 
                if expiry_date: days_left = (expiry_date - date.today()).days
            except: pass

        if is_active and 0 <= days_left <= 7:
            st.warning(f"⚠️ Subscription Alert: Your access expires in {days_left} days.")

        if not is_active or days_left < 0:
            # THE 3-STEP PAYMENT REDIRECT
            st.error("## ⛔ Access Restricted")
            st.info("Your subscription is inactive. Please use one of the payment methods below to restore access.")
            
            p1, p2 = st.columns(2)
            with p1:
                st.markdown("#### 📱 Option 1: MPESA")
                st.write(f"**Name:** {st.secrets['payments']['mpesa_name']}")
                st.write(f"**Number:** {st.secrets['payments']['mpesa_number']}")
                st.caption("Please send via Send Money.")
            
            with p2:
                st.markdown("#### 🏦 Option 2: Bank Transfer")
                st.write(f"**Bank:** {st.secrets['bank']['name']}")
                st.write(f"**Account Name:** {st.secrets['bank']['account_name']}")
                st.write(f"**Account No:** `{st.secrets['bank']['account_number']}`")
               

            st.divider()
            
            st.markdown("#### 🏁 Final Step: Activation")
            st.write("Take a screenshot of your payment and send it to the admin for instant activation.")
            
            # Using the admin whatsapp from your secrets
            wa_num = st.secrets['admin']['whatsapp']
            wa_url = f"https://wa.me/{wa_num}?text=Hi%20Ashley,%20I've%20paid%20for%20my%20Flight%20Portal%20subscription.%20Email:%20{user_email}"
            st.link_button("📤 Send Screenshot via WhatsApp", wa_url)
            
            st.stop()

# --- 6. DASHBOARD ---
        st.header("Welcome back")

        # --- NEW: READINESS CALCULATION LOGIC ---
        def get_status_color(expiry_date):
            if expiry_date is None: return "⚪ Not Found", "gray"
            # Ensure we are comparing Timestamp to Timestamp
            expiry_date = pd.Timestamp(expiry_date)
            days_diff = (expiry_date - pd.Timestamp.now().normalize()).days
            if days_diff < 0: return f"❌ Expired ({expiry_date.strftime('%Y-%m-%d')})", "red"
            if days_diff <= 30: return f"⚠️ Warning: {days_diff} Days Left", "orange"
            return f"✅ Valid until {expiry_date.strftime('%Y-%m-%d')}", "green"

        # Auto-IRT Logic: Look for 'IRT' in the 'DUTY' column of your logbook
               # Auto-IRT Logic: Look for 'IRT' in the 'DUTY' column of your logbook
irt_expiry = None  # This starts at the very edge (no spaces)

if not df_raw.empty and 'DUTY' in df_raw.columns:
    # Everything below is indented by 4 spaces
    search_pattern = 'IRT|CAT I|CAT II'
    
    # This line MUST line up with search_pattern
    irt_records = df_raw[df_raw['DUTY'].astype(str).str.contains(search_pattern, case=False, na=False)]
    
    if not irt_records.empty:
        # This is inside a second "if", so it is indented by 8 spaces
        last_irt_date = pd.to_datetime(irt_records['LOGBOOK DATE']).max()
        irt_expiry = last_irt_date + pd.DateOffset(months=6)

        col_btn, col_tip = st.columns([0.2, 0.8])
        with col_btn:
            if st.button("🔄 Sync Data"):
                st.rerun() 

        with col_tip:
            st.caption("**Tip:** Click **Sync** after adding new flight Records to reload the page")

        st.divider() 

        tab1, tab2 = st.tabs(["🏠 My Dashboard", "📖 Full Logbook"])
        
        with tab1:
            # --- NEW: PILOT READINESS INTERFACE ---
            st.subheader("🛡️ Pilot Readiness & Validity")
            r_col1, r_col2, r_col3 = st.columns(3)
            
        with r_col1:
            st.markdown("**CoFC (1 Year Validity)**")
            coc_ac = st.selectbox("Aircraft Type", ["C145A", "Y12 II"], key="coc_ac_input")
            
            # This logic ensures the date doesn't reset to today
            if 'coc_date' not in st.session_state:
                st.session_state.coc_date = date.today()
            
            last_coc = st.date_input("Date of Last CoC", value=st.session_state.coc_date, key="coc_sel_v3")
            st.session_state.coc_date = last_coc
            
            coc_exp = pd.Timestamp(last_coc) + pd.DateOffset(years=1)
            status, color = get_status_color(coc_exp)
            st.markdown(f":{color}[{status}]")

        with r_col2:
            st.markdown("**Medical (1 Year Validity)**")
            st.write("") 
            
            # This logic ensures the date doesn't reset to today
            if 'med_date' not in st.session_state:
                st.session_state.med_date = date.today()
            
            last_med = st.date_input("Date of Last Medical", value=st.session_state.med_date, key="med_sel_v3")
            st.session_state.med_date = last_med
            
            med_exp = pd.Timestamp(last_med) + pd.DateOffset(years=1)
            status, color = get_status_color(med_exp)
            st.markdown(f":{color}[{status}]")

        with r_col3:
            st.markdown("**IRT (6 Months Validity)**")
            st.caption("*(Auto-detected from Logbook 'Duty')*")
            
            # This follows your rule: looks for 'IRT' and 'CAT I & II'
            status, color = get_status_color(irt_expiry)
            st.markdown(f"**Current Status:**\n\n:{color}[{status}]")

        st.divider()

        # ---  CAREER TOTALS ---
        st.subheader("📊 Career Totals")
        c1, c2, c3 = st.columns(3)
        
        with c1:
            st.metric("Grand Total Hours", universal_formatter(safe_sum(df_raw, 'FLIGHT TIME Totals'), 'FLIGHT TIME Totals'))
        
        with c2:
            st.metric("Total Landings", universal_formatter(safe_sum(df_raw, 'LDGS'), 'LDGS'))
        
        with c3:
            st.metric("Total Instr. Flying", universal_formatter(safe_sum(df_raw, 'Total Instrument'), 'Total Instrument'))
        st.divider()

            # --- 7. MONTHLY NARRATIVE SEGMENT ---
        st.header("📅 Monthly Narrative & Summary")
        
        # 1. Pilot Info Inputs
        mcol1, mcol2, mcol3 = st.columns(3)
        with mcol1: m_rank = st.text_input("Rank (Narrative)", value=u.get("Rank", ""))
        with mcol2: m_name = st.text_input("Full Name (Narrative)", value=u.get("Name", ""))
        with mcol3: m_num = st.text_input("Service Number (Narrative)", value=u.get("Service Number", ""))

        # 2. Month/Year Selection
        ncol1, ncol2 = st.columns(2)
        month_list = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
        with ncol1: sel_month = st.selectbox("Select Month", month_list, index=pd.Timestamp.now().month - 1)
        with ncol2: sel_year = st.number_input("Select Year", min_value=2020, max_value=2030, value=2025)

        search_string = f"{sel_month} {int(sel_year)}"
        
        if 'df_stats' in locals() and not df_stats.empty:
            # Clean column names to prevent hidden space issues
            df_stats.columns = [str(c).strip() for c in df_stats.columns]
            
            # Helper to find the Month row
            month_cols = [c for c in df_stats.columns if 'month' in c.lower()]
            target_col = month_cols[0] if month_cols else df_stats.columns[0]
            narrative_row = df_stats[df_stats[target_col].astype(str) == search_string]

            if not narrative_row.empty:
                data = narrative_row.iloc[0]
                st.success(f"✅ Data found for {search_string}")
                
                # --- IMPROVED BULLETPROOF FORMATTER ---
                def final_time_converter(keyword):
                    cols = [c for c in df_stats.columns if keyword.lower() in c.lower()]
                    if not cols: return "0:00"
                    
                    val = data.get(cols[0], 0)
                    # Handle empty/NaN values
                    if pd.isna(val) or val == 0 or val == "": return "0:00"
                    
                    try:
                        # Convert to number to handle seconds (54300) or decimals (15.1)
                        num_val = float(val)
                        if num_val > 500: # It's raw seconds
                            h = int(num_val // 3600)
                            m = int((num_val % 3600) // 60)
                            return f"{h}:{m:02d}"
                        else: # It's decimal hours
                            h = int(num_val)
                            m = int(round((num_val - h) * 60))
                            return f"{h}:{m:02d}"
                    except (ValueError, TypeError):
                        # If it's already a string like "15:05", return it
                        return str(val)

                # --- FETCHING DATA ---
                # We use very specific keywords to match your Screenshot 305/306
                t_total  = final_time_converter('Total Flight Time')
                t_day    = final_time_converter('Total Day') 
                t_night  = final_time_converter('Total Night')
                t_actual = final_time_converter('Actual')
                t_sim    = final_time_converter('Sim')
                
                # Landings (A simple number)
                ldg_cols = [c for c in df_stats.columns if 'ldgs' in c.lower()]
                v_ldgs = int(data.get(ldg_cols[0], 0)) if ldg_cols else 0

                # --- UPDATED DISPLAY LAYOUT ---
                r1c1, r1c2, r1c3 = st.columns(3)
                r1c1.metric("Total Flight Time", t_total)
                r1c2.metric("Day Total", t_day)
                r1c3.metric("Night Total", t_night)

                r2c1, r2c2, r2c3 = st.columns(3)
                r2c1.metric("Instrument Actual", t_actual)
                r2c2.metric("Instrument Sim", t_sim)
                r2c3.metric("Landings", v_ldgs)

                st.divider()
                st.subheader("Monthly Narrative Text")
                # Look for the narrative text column specifically
                narr_col = [c for c in df_stats.columns if 'narrative' in c.lower()]
                narr_text = data.get(narr_col[0], "No text found.") if narr_col else "Narrative column not found."
                st.info(narr_text)
                
                # --- WORD DOCUMENT DOWNLOAD ---
                if st.button("📥 Download Official Monthly Narrative"):
                    from docx import Document
                    doc = Document()
                    doc.add_heading(f'MONTHLY FLYING NARRATIVE - {sel_month.upper()} {int(sel_year)}', 0).alignment = 1
                    
                    p = doc.add_paragraph()
                    p.add_run(f"RANK: {m_rank}\nNAME: {m_name}\nNUMBER: {m_num}").bold = True

                    st_table = doc.add_table(rows=2, cols=6)
                    st_table.style = 'Table Grid'
                    cols_headers = ["Total", "Day", "Night", "Actual", "Sim", "LDGS"]
                    vals_row = [t_total, t_day, t_night, t_actual, t_sim, str(v_ldgs)]
                    
                    for i, txt in enumerate(cols_headers):
                        st_table.rows[0].cells[i].text = txt
                        st_table.rows[1].cells[i].text = vals_row[i]

                    doc.add_heading('Narrative Summary', level=2)
                    doc.add_paragraph(str(narr_text))
                    
                    bio = BytesIO(); doc.save(bio)
                    st.download_button("Save Narrative", bio.getvalue(), f"Narrative_{search_string}.docx")
            else:
                st.warning(f"No records for {search_string} found in the Stats Table.")
        else:
            st.error("The Stats Table (tblNzncflTuufPioz) is empty or not loading.")

            # --- 7. FORM GENERATOR ---
        # 1. Convert the logbook date column to actual Timestamps
        df_raw['LOGBOOK DATE'] = pd.to_datetime(df_raw['LOGBOOK DATE'])

# 2. Calculate the 6-month cutoff as a Timestamp (NOT a .date())
        six_months_ago = pd.Timestamp.now() - pd.Timedelta(days=180)

# 3. Perform the comparison (Timestamp vs Timestamp)
        recent_df = df_raw[df_raw['LOGBOOK DATE'] >= six_months_ago]

        # --- 2. THE DOWNLOAD FUNCTION ---
        def generate_word_doc(form_type, p_rank, p_name, p_num, df_raw, recent_df, ac_type=None):
            doc = Document()
            if form_type == "CAT":
                doc.add_heading('ANNEX D - KENYA AIR FORCE CAT FORM P1', 0).alignment = 1
            else:
                doc.add_heading(f"PILOT'S INSTRUMENT RATING SCHEME - {ac_type}", 0).alignment = 1
                
            doc.add_paragraph(f"RANK: {p_rank}    NAME: {p_name}    NUMBER: {p_num}")

            # --- TABLE: FLYING HOURS (GENERAL) ---
            doc.add_heading('Flying Hours (General):', level=1)
            gen_map = [
                ("a. First Pilot (Turbo prop)", "1st Pilot (TP)"),
                ("b. Instrument Flying Actual (TP)", "Instr. Actual (TP)"),
                ("c. Instrument Flying Simulated (TP)", "Instr. Sim (TP)"),
                ("d. First Pilot (Piston)", "1st Pilot (Piston)"),
                ("e. Instrument Flying Actual (Piston)", "Instr. Actual (Piston)"),
                ("f. Instrument Flying Simulated (Piston)", "Instr. Sim (Piston)"),
                ("g. No. of Instrument Approaches (Total)", "I/F APPROACHES NO.")
            ]
            gt = doc.add_table(rows=1, cols=3); gt.style = 'Table Grid'
            gt.rows[0].cells[0].text = "Category"
            gt.rows[0].cells[1].text = "TOTAL"
            gt.rows[0].cells[2].text = "LAST 6 MONTHS"

            for label, col in gen_map:
                row = gt.add_row().cells
                row[0].text = label
                row[1].text = universal_formatter(safe_sum(df_raw, col), "Time")
                row[2].text = universal_formatter(safe_sum(recent_df, col), "Time")

            if form_type == "CAT":
                # --- CAT TABLE 1: FLYING TIMES (DAY & NIGHT) ---
                doc.add_heading('1. Flying Times Breakdown', level=1)
                ft = doc.add_table(rows=3, cols=7); ft.style = 'Table Grid'
                headers = ["Type", "Total", "Twin", "Sim", "Actual", "Let Down", "Ldg"]
                for i, h in enumerate(headers): ft.rows[0].cells[i].text = h
                
                for idx, label in enumerate(["DAY", "NIGHT"], 1):
                    r = ft.rows[idx].cells
                    suffix = " Day" if label == "DAY" else " Night"
                    r[0].text = label
                    r[1].text = universal_formatter(safe_sum(df_raw, f"{label.title()} Total"), "Time")
                    r[2].text = universal_formatter(get_twin_logic(df_raw, f"{label.title()} Total"), "Time")
                    r[3].text = universal_formatter(safe_sum(df_raw, f"Instr. Flying Sim{suffix}"), "Time")
                    r[4].text = universal_formatter(safe_sum(df_raw, f"Instr. Flying Actual{suffix}"), "Time")
                    r[5].text = str(int(safe_sum(recent_df, "I/F APPROACHES NO.")))
                    r[6].text = str(int(safe_sum(recent_df, "Landings (6 Months)")))


            elif form_type == "IRM":
                doc.add_heading(f'2. Flying Hours on Type: {ac_type}', level=1)
                it = doc.add_table(rows=3, cols=3); it.style = 'Table Grid'
                it.rows[0].cells[0].text = "Description"; it.rows[0].cells[1].text = "TOTAL"; it.rows[0].cells[2].text = "LAST 6 MONTHS"
                target_col = f"Total {ac_type.replace(' ', '')}"
                it.rows[1].cells[0].text = "b. Total 1st Pilot"
                it.rows[1].cells[1].text = universal_formatter(safe_sum(df_raw, target_col), "Time")
                it.rows[1].cells[2].text = universal_formatter(safe_sum(recent_df, target_col), "Time")
                it.rows[2].cells[0].text = "c. No. of Instrument Approaches"
                it.rows[2].cells[1].text = str(int(safe_sum(df_raw, "I/F APPROACHES NO.")))
                it.rows[2].cells[2].text = str(int(safe_sum(recent_df, "I/F APPROACHES NO.")))

            bio = BytesIO(); doc.save(bio)
            return bio.getvalue()

        # --- 3. STREAMLIT INTERFACE ---
        st.subheader("📝 Official Form Portal")
        fcol1, fcol2, fcol3 = st.columns(3)
        # These variable names (p_rank, p_name, p_num) MUST match the function call below
        with fcol1: p_rank = st.text_input("Rank", value=u.get("Rank", ""))
        with fcol2: p_name = st.text_input("Full Name", value=u.get("Name", ""))
        with fcol3: p_num = st.text_input("Service Number", value=u.get("Service Number", ""))

        # ENSURE BOTH ARE IN THIS LIST
        form_choice = st.selectbox("Select Form Type", ["Blank Cat Form", "IR Masila Form"])

        # --- 2. DATA PREP (6-Month Filter) ---
        recent_df = df_raw[df_raw['LOGBOOK DATE'] >= (pd.Timestamp.now() - pd.Timedelta(days=180))]

        # --- 3. FORM REDERING LOGIC ---
        if form_choice == "Blank Cat Form":
            st.info("Generating CAT 55 with Day/Night Rows.")
            # Trigger download function
            cat_file = generate_word_doc("CAT", p_rank, p_name, p_num, df_raw, recent_df)
            st.download_button("📥 Download CAT 55 (.docx)", cat_file, f"CAT55_{p_name}.docx")
        
        elif form_choice == "IR Masila Form":
            # This section only appears when "IR Masila Form" is selected in the dropdown
            st.write("### IR Masila Configuration")
            ac_type = st.selectbox("Select Aircraft for IRM", ["C145A", "Y12 II"])
            
            # Generate IRM specific doc
            irm_file = generate_word_doc("IRM", p_rank, p_name, p_num, df_raw, recent_df, ac_type=ac_type)
            st.download_button(f"📥 Download IR Masila ({ac_type})", irm_file, f"IRM_{ac_type}_{p_name}.docx")

        with tab2:
            st.subheader("📖 Complete Flying Record (Latest First)")
            
            if not df_raw.empty:
                # --- 1. SEARCH BAR ---
                search_query = st.text_input("🔍 Search Logbook (Type to filter by Aircraft, Route, or Date)", key="logbook_search").lower()

                # --- DATA PROCESSING & FORMATTING ---
                df_raw.columns = df_raw.columns.str.strip()
                display_df = df_raw.copy()
                
                target_date_col = "LOGBOOK DATE"
                
                if target_date_col in display_df.columns:
                    # Format the date and rename it for the display
                    display_df[target_date_col] = pd.to_datetime(display_df[target_date_col]).dt.strftime('%Y-%m-%d')
                    display_df = display_df.rename(columns={target_date_col: "DATE"})
                
                # Formatting other columns
                for col in display_df.columns:
                    col_lower = col.lower()
                    if col == "DATE":
                        continue
                    # Apply the HH:MM formatter to time/flying columns
                    if any(word in col_lower for word in ["time", "flying", "total", "dual", "tp", "piston", "instr", "ldg", "approach"]):
                        display_df[col] = display_df[col].apply(lambda x: universal_formatter(x, col))

                # --- SEARCH FILTER ---
                if search_query:
                    mask = display_df.astype(str).apply(lambda x: x.str.contains(search_query, case=False)).any(axis=1)
                    display_df = display_df[mask]

                # --- DISPLAY ---
                if not display_df.empty:
                    # Ensure DATE is the first column if it exists
                    if "DATE" in display_df.columns:
                        cols = ["DATE"] + [c for c in display_df.columns if c != "DATE"]
                        display_df = display_df[cols]
                    
                    # This now only renders inside Tab 2
                    st.dataframe(display_df, use_container_width=True, hide_index=True)
                    st.caption("💡 **Tip:** Hover over the table and click the two diagonal squares (top right) for a clear view.")
                else:
                    st.info("No records found.")
            else:
                st.info("No flight records available to display.")

else:
    # This is for the very bottom of the script
    if not user_email:

        st.info("### 🛫 Please login in the sidebar to access your flight portal.")



